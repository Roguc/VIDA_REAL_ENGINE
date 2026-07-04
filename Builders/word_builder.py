from pathlib import Path

from docx import Document

from Models.sistema_dia import SistemaDia


class WordBuilder:
    def __init__(self, root_dir: Path):
        self.root_dir = Path(root_dir)
        self.out_dir = self.root_dir / "Salidas" / "Word"
        self.out_dir.mkdir(parents=True, exist_ok=True)

    def build(self, s: SistemaDia) -> Path:
        doc = Document()
        doc.add_heading(f"Vida Real Engine V5 — {s.fecha_iso}", level=1)
        doc.add_paragraph("Buenos días, Robinson. Este es el respaldo documental del Centro de Comando.")
        doc.add_heading("Resumen", level=2)
        for k, v in s.resumen().items():
            doc.add_paragraph(f"{k}: {v}")
        doc.add_heading("Cronograma", level=2)
        for b in s.cronograma:
            doc.add_paragraph(f"{b['hora_inicio']}–{b['hora_fin']} | {b['titulo']}")
            doc.add_paragraph(b['actividad'])
            doc.add_paragraph(f"Registro: {b['registro']}")
        out = self.out_dir / f"VIDA_REAL_ENGINE_V5_{s.fecha_iso}.docx"
        doc.save(out)
        return out
