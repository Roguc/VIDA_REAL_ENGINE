from datetime import datetime, date
from openpyxl import load_workbook
from docx import Document

from Config.config import EXCEL_DIR, WORD_DIR, LOGS_DIR, HOJAS_MOTOR
from Scripts.utilidades import limpiar_texto, buscar_excel, escribir_log, hoy, nombre_dia_es
from Motores.motor_serpat import obtener_turno_fecha
from Motores.motor_fuentes import leer_pendientes, leer_vision_board, leer_desarrollo_personal
from Motores.motor_universidad import construir_bloque_universidad
from Motores.motor_finanzas import construir_bloque_finanzas
from Motores.motor_salud import construir_bloque_salud
from Motores.motor_empresas import construir_bloque_empresas
from Motores.motor_ancla import construir_bloque_ancla
from Motores.motor_continuidad import capital_del_bloque, adaptar_por_energia

def abrir_excel():
    ruta_excel = buscar_excel(EXCEL_DIR)
    wb = load_workbook(ruta_excel, data_only=True)
    return ruta_excel, wb

def verificar_hojas_motor(wb):
    faltantes = [h for h in HOJAS_MOTOR if h not in wb.sheetnames]
    if faltantes:
        raise ValueError(f"Faltan hojas motor: {faltantes}")

def leer_parametros_base(wb):
    datos = {}
    if "00_GENERADOR_DIA" not in wb.sheetnames:
        return datos
    ws = wb["00_GENERADOR_DIA"]
    for fila in ws.iter_rows(min_row=1, max_row=45, values_only=True):
        if not fila or fila[0] is None:
            continue
        clave = limpiar_texto(fila[0])
        valor = fila[1] if len(fila) > 1 else ""
        if clave:
            datos[clave] = valor
    return datos

def construir_dia(wb, fecha_objetivo=None):
    fecha = fecha_objetivo or hoy()
    datos = leer_parametros_base(wb)
    turno = obtener_turno_fecha(wb, fecha)
    desarrollo = leer_desarrollo_personal(wb)

    dia = {
        "fecha": fecha.strftime("%Y-%m-%d"),
        "fecha_obj": fecha,
        "fecha_larga": f"{nombre_dia_es(fecha)} {fecha.strftime('%d-%m-%Y')}",
        "tipo_dia": limpiar_texto(datos.get("Tipo día")) or "Día operativo",
        "turno_serpat": limpiar_texto(datos.get("Turno SERPAT")),
        "semana_academica": limpiar_texto(datos.get("Semana académica")),
        "prioridad_academica": limpiar_texto(datos.get("Prioridad académica")),
        "evaluacion_critica": limpiar_texto(datos.get("Evaluación crítica")),
        "prioridad_empresa": limpiar_texto(datos.get("Prioridad empresa")),
        "nivel_energia": limpiar_texto(datos.get("Nivel energía")) or "Pendiente",
        "alerta_salud": limpiar_texto(datos.get("Presión / alerta salud")) or "Pendiente",
        "cierre_mes": limpiar_texto(datos.get("Cierre de mes")),
        "desarrollo": desarrollo,
        "pendientes": leer_pendientes(wb),
        "vision": leer_vision_board(wb),
        "serpat": turno,
    }

    if turno:
        dia["tipo_dia"] = turno["tipo_dia"]
        dia["turno_serpat"] = turno["turno_serpat"]
        dia["serpat_observaciones"] = turno["observaciones"]
        dia["serpat_estado"] = turno["estado"]
        dia["serpat_ingreso"] = turno["ingreso"]
        dia["serpat_salida"] = turno["salida"]

    return dia

def armar_bloque(hora_inicio, hora_fin, bloque):
    area = bloque.get("area", "")
    return {
        "hora_inicio": hora_inicio,
        "hora_fin": hora_fin,
        "area": area,
        "capital": capital_del_bloque(area),
        "titulo": bloque.get("titulo", ""),
        "objetivo": bloque.get("objetivo", ""),
        "actividad": bloque.get("actividad", ""),
        "registro": bloque.get("registro", ""),
        "evidencia": bloque.get("evidencia", ""),
        "seguimiento": bloque.get("seguimiento", ""),
    }

def construir_cronograma(dia):
    tipo = dia.get("tipo_dia", "").lower()

    if "noche" in tipo:
        bloques = [
            armar_bloque("08:00","08:20",{"area":"SERPAT","titulo":"SALIDA / ENTREGA DEL TURNO","objetivo":"Cerrar correctamente el turno.","actividad":"Entregar novedades, verificar observaciones pendientes y salir del recinto.","registro":"No corresponde","evidencia":"✓ Turno entregado.","seguimiento":"Registrar novedades relevantes."}),
            armar_bloque("08:20","09:10",construir_bloque_salud(dia) | {"titulo":"TRASLADO · LLEGADA A CASA · DESCANSO"}),
            armar_bloque("09:10","14:30",{"area":"Salud","titulo":"SUEÑO PRINCIPAL","objetivo":"Recuperar energía.","actividad":"Dormir, teléfono en silencio, dormitorio oscuro y sin interrupciones.","registro":"H02_Registro_Salud al despertar","evidencia":"✓ Descanso completado.","seguimiento":"Registrar horas de sueño."}),
            armar_bloque("14:30","14:40",{"area":"Ancla Mental","titulo":"DESPERTAR · IDENTIDAD · VISUALIZACIÓN","objetivo":"Activar Robinson 2042.","actividad":"Tomar agua. Respirar. Visualizar panel de visión: oficina, holding, salud, familia, patrimonio y libertad.","registro":"Ancla Mental; 21_KPI_Diario","evidencia":"✓ Visualización realizada.","seguimiento":"Registrar estado mental inicial 1-10."}),
            armar_bloque("14:40","15:00",construir_bloque_salud(dia)),
            armar_bloque("15:00","15:30",{"area":"Salud","titulo":"ALIMENTACIÓN · HIDRATACIÓN","objetivo":"Construir capital biológico.","actividad":"Almorzar, hidratarse, evitar azúcar/gasto impulsivo y evaluar energía posterior.","registro":"H02_Registro_Salud; 12_Gastos si hubo compra","evidencia":"✓ Alimentación registrada.","seguimiento":"Registrar comida real."}),
            armar_bloque("15:30","15:50",{"area":"Sistema","titulo":"APERTURA OFICIAL DEL SISTEMA","objetivo":"Conocer la realidad antes de ejecutar.","actividad":"Revisar dashboard, KPI, calendario laboral, calendario académico, alertas, finanzas, universidad, empresas y salud.","registro":"20_Registro_Diario; 21_KPI_Diario; 90_Alertas_Correcciones","evidencia":"✓ Sistema aperturado.","seguimiento":"Confirmar prioridad única."}),
            armar_bloque("15:50","16:05",{"area":"Universidad","titulo":"ALERTA ACADÉMICA · EVALUACIONES","objetivo":"Evitar perder pruebas o entregas.","actividad":f"Revisar Canvas y 62_T2_Evaluaciones. Evaluación crítica: {dia.get('evaluacion_critica')}.","registro":"62_T2_Evaluaciones; 66_T2_Calendario","evidencia":"✓ Plazos confirmados.","seguimiento":"Anotar enviado / pendiente / corregir / estudiar."}),
            armar_bloque("16:05","17:20",construir_bloque_universidad(dia)),
            armar_bloque("17:20","17:50",{"area":"Universidad","titulo":"COMUNICACIÓN Y ARGUMENTACIÓN","objetivo":"Construir capital comunicacional.","actividad":"Revisar pauta, rúbrica, tesis, argumentos, coherencia, conectores, ortografía y formato.","registro":"61_T2_Ramos; 62_T2_Evaluaciones; 63_T2_Bloques","evidencia":"✓ Rúbrica validada.","seguimiento":"Registrar pendiente exacto."}),
            armar_bloque("17:50","18:10",{"area":"Universidad","titulo":"CÁLCULO II · CONTINUIDAD MÍNIMA","objetivo":"Evitar que Cálculo siga en cero.","actividad":"Resolver mínimo 1 ejercicio si estás cansado, ideal 3. Registrar error exacto.","registro":"61_T2_Ramos; 63_T2_Bloques; 64_T2_Errores","evidencia":"✓ Ejercicio desarrollado.","seguimiento":"Registrar ejercicio exacto."}),
            armar_bloque("18:10","18:25",{"area":"Universidad","titulo":"FÍSICA MECÁNICA · CONTINUIDAD MÍNIMA","objetivo":"Evitar acumulación.","actividad":"Resolver mínimo 1 ejercicio conceptual o numérico. Registrar duda concreta.","registro":"61_T2_Ramos; 63_T2_Bloques; 64_T2_Errores","evidencia":"✓ Física tocada.","seguimiento":"Registrar fórmula o concepto no entendido."}),
            armar_bloque("18:25","18:45",construir_bloque_empresas(dia)),
            armar_bloque("18:45","19:00",construir_bloque_finanzas(dia)),
            armar_bloque("19:00","19:10",{"area":"Desarrollo personal","titulo":"LECTURA ESTRATÉGICA","objetivo":"Construir capital cognitivo.","actividad":f"Leer 10 minutos: {dia['desarrollo'].get('libro') or 'lectura estratégica'} capítulo {dia['desarrollo'].get('capitulo') or 'pendiente'}. Registrar una idea aplicable.","registro":"Ancla Mental; Registro Diario","evidencia":"✓ Lectura realizada.","seguimiento":"Registrar frase que cambie conducta."}),
            armar_bloque("19:10","19:20",{"area":"Imagen personal","titulo":"IMAGEN PERSONAL · PRESENCIA","objetivo":"Presentarte como Robinson en ascenso.","actividad":"Uniforme limpio, higiene, postura, mochila ordenada, credencial, agua y colación.","registro":"21_KPI_Diario","evidencia":"✓ Presencia cuidada.","seguimiento":"Registrar atraso/desorden si ocurrió."}),
            armar_bloque("19:20","20:20",{"area":"SERPAT","titulo":"TRASLADO A SERPAT","objetivo":"Llegar estable y puntual.","actividad":"Trasladarse, evitar redes, preparar mentalmente el turno y entrar con calma.","registro":"No corresponde","evidencia":"✓ Llegada anticipada.","seguimiento":"Prioridad seguridad."}),
            armar_bloque("20:20","21:00",{"area":"SERPAT","titulo":"PREPARACIÓN OPERACIONAL DEL TURNO","objetivo":"Comenzar el turno preparado.","actividad":"Revisar servicio, protocolos, equipos, llaves, radio y libro de novedades.","registro":"No corresponde","evidencia":"✓ Puesto operativo.","seguimiento":"Aclarar dudas antes del inicio."}),
            armar_bloque("21:00","08:00",{"area":"SERPAT","titulo":"SERPAT — TURNO NOCTURNO","objetivo":"Construir capital reputacional.","actividad":"Cumplir responsabilidades, mantener atención, hidratarse, consumir colación y registrar mentalmente novedades.","registro":"No corresponde","evidencia":"✓ Turno ejecutado correctamente.","seguimiento":"Novedades para entrega."}),
        ]
    else:
        bloques = [
            armar_bloque("05:30","05:45",{"area":"Ancla Mental","titulo":"DESPERTAR · IDENTIDAD · VISUALIZACIÓN","objetivo":"Activar Robinson 2042.","actividad":"Abrir los ojos. Tomar agua. Respirar. Visualizar panel de visión. Repetir: hoy no improviso, sigo el sistema.","registro":"Ancla Mental; 21_KPI_Diario","evidencia":"✓ Visualización realizada.","seguimiento":"Registrar estado mental inicial 1-10."}),
            armar_bloque("05:45","06:15",construir_bloque_salud(dia)),
            armar_bloque("06:15","06:45",{"area":"Imagen personal","titulo":"HIGIENE · VESTIMENTA · PRESENCIA","objetivo":"Construir presencia profesional diaria.","actividad":"Ducha, higiene, ropa ordenada, postura, revisar bolso y preparar salida.","registro":"21_KPI_Diario","evidencia":"✓ Presencia cuidada.","seguimiento":"Registrar atraso/desorden."}),
            armar_bloque("06:45","07:10",{"area":"Salud","titulo":"DESAYUNO · MEDICAMENTO · AGUA","objetivo":"Preparar energía estable.","actividad":"Desayunar, hidratarse, tomar medicamento si corresponde y registrar presión si aplica.","registro":"H02_Registro_Salud; H03_Presión_Log","evidencia":"✓ Alimentación/medicación registrada.","seguimiento":"Registrar comida real."}),
            armar_bloque("07:10","07:30",{"area":"Sistema","titulo":"APERTURA DEL SISTEMA","objetivo":"Leer realidad antes de ejecutar.","actividad":"Revisar Excel C31, calendario, evaluaciones, salud, finanzas y prioridades.","registro":"20_Registro_Diario; 21_KPI_Diario; 90_Alertas_Correcciones","evidencia":"✓ Día aperturado.","seguimiento":"Definir prioridad única."}),
            armar_bloque("07:30","08:30",construir_bloque_universidad(dia)),
            armar_bloque("08:30","09:00",construir_bloque_finanzas(dia)),
            armar_bloque("09:00","09:30",construir_bloque_empresas(dia)),
            armar_bloque("20:30","21:00",construir_bloque_ancla(dia)),
        ]

    return [adaptar_por_energia(dia, b) for b in bloques]

def generar_texto(dia, cronograma):
    lineas = [f"VIDA REAL ENGINE — {dia.get('fecha')}", "", "Buenos días, Robinson.", "Este es tu plan operativo del día.", ""]
    for b in cronograma:
        lineas += [f"{b['hora_inicio']}–{b['hora_fin']}", b["titulo"], "Objetivo", b["objetivo"], "Actividad", b["actividad"], "Registro", b["registro"], ""]
    return "\n".join(lineas)

def guardar_word(dia, texto):
    doc = Document()
    doc.add_heading(f"Vida Real Engine — {dia.get('fecha')}", level=1)
    for linea in texto.split("\n"):
        doc.add_paragraph(linea)
    ruta = WORD_DIR / f"VIDA_REAL_ENGINE_{dia.get('fecha')}.docx"
    doc.save(ruta)
    return ruta

def ejecutar_motor_vida_real(fecha_objetivo=None):
    ruta_log = LOGS_DIR / "vida_real_engine.log"
    ruta_excel, wb = abrir_excel()
    verificar_hojas_motor(wb)
    dia = construir_dia(wb, fecha_objetivo)
    cronograma = construir_cronograma(dia)
    texto = generar_texto(dia, cronograma)
    ruta_word = guardar_word(dia, texto)
    escribir_log(ruta_log, f"V4 generado: {ruta_word.name}")

    print("✅ VIDA REAL ENGINE V4 EJECUTADO")
    print(f"Excel leído: {ruta_excel.name}")
    print(f"Fecha automática: {dia.get('fecha')}")
    print(f"Tipo día: {dia.get('tipo_dia')}")
    print(f"Turno SERPAT: {dia.get('turno_serpat')}")
    print(f"Bloques generados: {len(cronograma)}")
    print(f"Word respaldo: {ruta_word}")
    print()

    return {"dia": dia, "cronograma": cronograma, "texto": texto, "word": ruta_word, "excel": ruta_excel}
